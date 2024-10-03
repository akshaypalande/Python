import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Function to add background color to cells (optional)
def set_cell_background(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_receipt(waiter_name, table_name, date, time, items):
    # Create a new document
    doc = docx.Document()

    # Set the page size (25 cm height, 8 cm width)
    section = doc.sections[0]
    section.page_height = Cm(25)
    section.page_width = Cm(8)

    # Set the margins
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)

    # Set the default font style and size for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'  # Font similar to receipt print
    font.size = Pt(9)  # Reduced font size slightly

    # Add the restaurant name in center
    paragraph = doc.add_paragraph("Akshay Entertainment Pvt")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add restaurant address
    address = "Akshay Restaurant\nGround and First Floor, Site 16/1, P&T Layout,\nDr Shivaram Karanth Nagar, Bengaluru, Karnataka 560077"
    paragraph = doc.add_paragraph(address)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add bill details
    details = f"DATE: {date}\t\tTIME: {time}\nBill No: 409\t\tTable: {table_name}\nWaiter: {waiter_name}"
    paragraph = doc.add_paragraph(details)
    
    # Add some space
    doc.add_paragraph(" ")

    # Create table for itemized billing
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Add header cells (Items, Qty, Price)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Items'
    hdr_cells[1].text = 'Qty'
    hdr_cells[2].text = 'Price'
    
    total_amount = 0
    # Add items to the table
    for item, qty, unit_price in items:
        total_price = qty * unit_price  # Calculate total price for the item
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = str(qty)
        row_cells[2].text = f"{total_price:.2f}"  # Total price for the quantity
        total_amount += total_price  # Update total amount
    
    # Add total items
    doc.add_paragraph(f"\nTotal No. Items: {len(items)}")
    doc.add_paragraph(f"Sub Total: {total_amount:.2f}")
    
    # Calculate taxes and charges
    cgst = total_amount * 0.025
    sgst = total_amount * 0.025
    service_charge = total_amount * 0.06
    total_with_charges = total_amount + cgst + sgst + service_charge

    # Add tax details and total
    doc.add_paragraph(f"Central GST @2.5%: {cgst:.2f}")
    doc.add_paragraph(f"State GST @2.5%: {sgst:.2f}")
    doc.add_paragraph(f"Service Charge @6%: {service_charge:.2f}")
    doc.add_paragraph(f"Round Off: 0.04")
    doc.add_paragraph(f"Total: {total_with_charges + 0.04:.2f}")

    # Add total in words (this is just a sample, you may use a library to convert to words)
    doc.add_paragraph("\nSeven Hundred and Forty-Six")
    
    # Add payment method details
    doc.add_paragraph("\nSettlements:")
    doc.add_paragraph("CREDIT CARD\nVISA\tC.No.: XXXX1545")
    
    # Add GSTIN and thank you message
    doc.add_paragraph("\nGSTIN: 1234akshay12345")
    doc.add_paragraph("FSSAI: ")
    doc.add_paragraph("Thank You Visit Again.")

    # Save the document
    doc.save("bill.docx")

# Example usage
if __name__ == "__main__":
    # Inputs from the user
    waiter_name = input("Enter the waiter's name: ")
    table_name = input("Enter the table name: ")
    date = input("Enter the date (dd-mm-yy): ")
    time = input("Enter the time (hh:mm): ")

    # Items with quantity and price entered by user
    items = []
    while True:
        item_name = input("Enter the dish name (or 'done' to finish): ")
        if item_name.lower() == 'done':
            break
        quantity = int(input(f"Enter the quantity for {item_name}: "))
        unit_price = float(input(f"Enter the unit price for {item_name}: "))  # Price for 1 unit
        items.append((item_name, quantity, unit_price))  # Add to the list
    
    # Generate the receipt
    create_receipt(waiter_name, table_name, date, time, items)
